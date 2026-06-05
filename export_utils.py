# export_utils.py
import json
import os
from datetime import datetime
import streamlit as st
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

def generate_markdown(gap_report, summaries, papers, topic):
    """Generate a Markdown string of the gap report."""
    n_papers = len([p for p in papers if p.get('readable')])
    n_themes = len(summaries)
    n_gaps = len(gap_report.get('gaps', []))
    md = f"# Research Gap Analysis Report\n\n"
    md += f"**Topic:** {topic}\n\n"
    md += f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
    md += f"**Papers analysed:** {n_papers} | **Themes found:** {n_themes} | **Gaps identified:** {n_gaps}\n\n"
    md += "## Themes\n\n"
    for i, s in enumerate(summaries):
        sparse = " *(sparse)*" if s.get('sparse') else ""
        md += f"### Theme {i}{sparse} ({s['percentage']}% of chunks)\n"
        md += f"{s['summary']}\n\n"
    md += "## Identified Research Gaps\n\n"
    for gap in gap_report['gaps']:
        md += f"### {gap['title']}\n"
        md += f"- **Priority:** {gap['priority']}\n"
        md += f"- **Type:** {gap['type']}\n"
        md += f"- **Evidence themes:** {', '.join([str(t) for t in gap['evidence_themes']])}\n"
        md += f"- **Description:** {gap['description']}\n"
        md += f"- **Suggested direction:** {gap['suggested_direction']}\n\n"
    return md

def export_markdown(gap_report, summaries, papers, topic):
    md = generate_markdown(gap_report, summaries, papers, topic)
    st.download_button("Download Markdown", md, file_name="gap_report.md", mime="text/markdown")

def export_pdf(gap_report, summaries, papers, topic):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    
    # -- Title --
    pdf.set_font("Helvetica", 'B', 18)
    pdf.cell(0, 12, "Research Gap Analysis Report", ln=True, align='C')
    pdf.ln(6)
    
    # -- Meta info --
    n_papers = len([p for p in papers if p.get('readable')])
    n_themes = len(summaries)
    n_gaps = len(gap_report.get('gaps', []))
    pdf.set_font("Helvetica", '', 11)
    pdf.cell(0, 7, f"Topic: {topic}", ln=True)
    pdf.cell(0, 7, f"Papers analysed: {n_papers} | Themes: {n_themes} | Gaps: {n_gaps}", ln=True)
    pdf.ln(5)
    
    # -- Themes --
    pdf.set_font("Helvetica", 'B', 14)
    pdf.cell(0, 10, "Themes", ln=True)
    pdf.ln(2)
    for i, s in enumerate(summaries):
        sparse = " (sparse)" if s.get('sparse') else ""
        pdf.set_font("Helvetica", 'B', 12)
        pdf.cell(0, 8, f"Theme {i}{sparse} ({s['percentage']}%)", ln=True)
        pdf.set_font("Helvetica", '', 11)
        pdf.multi_cell(0, 6, s['summary'])
        pdf.ln(4)
    
    # -- Gaps --
    pdf.add_page()
    pdf.set_font("Helvetica", 'B', 14)
    pdf.cell(0, 10, "Identified Research Gaps", ln=True)
    pdf.ln(2)
    for gap in gap_report['gaps']:
        # Title
        pdf.set_font("Helvetica", 'B', 12)
        pdf.multi_cell(0, 7, gap['title'])
        # Priority and type
        pdf.set_font("Helvetica", '', 11)
        pdf.cell(0, 7, f"Priority: {gap['priority']}     Type: {gap['type']}", ln=True)
        # Evidence themes
        evidence = ', '.join([str(t) for t in gap['evidence_themes']])
        pdf.cell(0, 7, f"Evidence themes: {evidence}", ln=True)
        pdf.ln(2)
        # Description
        pdf.set_font("Helvetica", '', 11)
        pdf.multi_cell(0, 6, f"Description: {gap['description']}")
        pdf.ln(2)
        # Suggested direction
        pdf.set_font("Helvetica", 'I', 11)
        pdf.multi_cell(0, 6, f"Suggested direction: {gap['suggested_direction']}")
        pdf.ln(6)   # extra space between gaps
    
    # Convert bytearray to bytes for Streamlit
    pdf_bytes = bytes(pdf.output())
    st.download_button("Download PDF", pdf_bytes, file_name="gap_report.pdf", mime="application/pdf")

def export_docx(gap_report, summaries, papers, topic):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    # Title
    title = doc.add_heading('Research Gap Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Info
    n_papers = len([p for p in papers if p.get('readable')])
    n_themes = len(summaries)
    n_gaps = len(gap_report.get('gaps', []))
    doc.add_paragraph(f"Topic: {topic}")
    doc.add_paragraph(f"Papers analysed: {n_papers} | Themes found: {n_themes} | Gaps identified: {n_gaps}")
    doc.add_paragraph("")
    # Themes
    doc.add_heading('Themes', level=1)
    for i, s in enumerate(summaries):
        sparse = " (sparse)" if s.get('sparse') else ""
        doc.add_heading(f"Theme {i}{sparse} ({s['percentage']}%)", level=2)
        doc.add_paragraph(s['summary'])
    # Gaps
    doc.add_heading('Identified Research Gaps', level=1)
    for gap in gap_report['gaps']:
        doc.add_heading(gap['title'], level=2)
        p = doc.add_paragraph()
        p.add_run('Priority: ').bold = True
        p.add_run(f"{gap['priority']}   ")
        p.add_run('Type: ').bold = True
        p.add_run(gap['type'])
        doc.add_paragraph(f"Evidence themes: {', '.join([str(t) for t in gap['evidence_themes']])}")
        doc.add_paragraph(f"Description: {gap['description']}")
        doc.add_paragraph(f"Suggested direction: {gap['suggested_direction']}")
        doc.add_paragraph("")
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button("Download DOCX", buffer, file_name="gap_report.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")